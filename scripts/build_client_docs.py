from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"
OUT.mkdir(exist_ok=True)

NAVY = "0E2238"
INK = "1B2733"
MUTED = "65717E"
GOLD = "B88950"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F4F6F9"
WHITE = "FFFFFF"
RED = "9B1C1C"
GREEN = "1F5D42"
BODY_FONT = "Calibri"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(value):
    return RGBColor.from_string(value)


def set_run(run, size=None, color=INK, bold=None, italic=None, font=BODY_FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa):
    assert sum(widths_dxa) == TABLE_WIDTH_DXA
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)


def configure_page(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def configure_styles(doc, preset):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6 if preset == "compact" else 8)
    normal.paragraph_format.line_spacing = 1.25 if preset == "compact" else 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT if preset == "compact" else WD_ALIGN_PARAGRAPH.JUSTIFY

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14 if preset == "compact" else 12, 7 if preset == "compact" else 6),
        "Heading 3": (12, DARK_BLUE, 10 if preset == "compact" else 8, 5 if preset == "compact" else 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def add_numbering(doc, preset):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    start_abs = max(existing_abs or [0]) + 1
    start_num = max(existing_num or [0]) + 1
    if preset == "compact":
        left, hanging, after, line = 540, 271, 80, 300
    else:
        left, hanging, after, line = 540, 279, 80, 290

    def make(fmt, text, abs_id, num_id):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), str(after))
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abs_id))
        num.append(abs_ref)
        numbering.append(num)
        return num_id

    return {
        "bullet": make("bullet", "•", start_abs, start_num),
        "number": make("decimal", "%1.", start_abs + 1, start_num + 1),
    }


def add_page_field(paragraph, field):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run(run, size=8.5, color=MUTED)


def set_header_footer(doc, left, right):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    set_run(p.add_run(left.upper()), size=8.5, color=NAVY, bold=True)
    set_run(p.add_run("\t" + right), size=8.5, color=MUTED)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D7DBE2")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(fp.add_run("REGOLA  |  "), size=8.5, color=MUTED)
    add_page_field(fp, "PAGE")
    set_run(fp.add_run(" / "), size=8.5, color=MUTED)
    add_page_field(fp, "NUMPAGES")


def add_title_block(doc, kicker, title, subtitle, meta):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(kicker.upper()), size=9, color=GOLD, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run(title), size=29, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    set_run(p.add_run(subtitle), size=13, color=MUTED)
    table = add_table(doc, [["ПОДГОТОВЛЕНО ДЛЯ", meta[0]], ["СТАТУС", meta[1]], ["ДАТА", meta[2]]], [2700, 6660], header=False)
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        set_cell_shading(row.cells[0], PALE_BLUE)
        for run in row.cells[0].paragraphs[0].runs:
            set_run(run, size=8.5, color=DARK_BLUE, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_editorial_cover(doc, kicker, title, subtitle, edition):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(105)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(kicker.upper()), size=9, color=GOLD, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run(title), size=31, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(38)
    set_run(p.add_run(subtitle), size=14, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(80)
    set_run(p.add_run("Премиальная фурнитура без визуального шума"), size=11, color=GOLD, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(edition), size=10, color=MUTED)


def add_paragraph(doc, text, bold_lead=None, italic=False, color=INK):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True, color=color)
        set_run(p.add_run(text[len(bold_lead):]), color=color, italic=italic)
    else:
        set_run(p.add_run(text), color=color, italic=italic)
    return p


def add_list(doc, items, num_id, preset):
    for item in items:
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num])
        p_pr.append(num_pr)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25 if preset == "compact" else 1.208
        set_run(p.add_run(item), size=11)


def add_callout(doc, label, text, kind="info"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.18
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE_GRAY)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), RED if kind == "warning" else GOLD)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    set_run(p.add_run(label.upper() + "  "), size=9.5, color=RED if kind == "warning" else DARK_BLUE, bold=True)
    set_run(p.add_run(text), size=10.5, color=INK)


def add_table(doc, rows, widths_dxa, header=True, font_size=9.5):
    table = doc.add_table(rows=len(rows), cols=len(widths_dxa))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    for r_idx, row_values in enumerate(rows):
        for c_idx, value in enumerate(row_values):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = str(value)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                set_run(run, size=font_size, color=INK, bold=header and r_idx == 0)
            if header and r_idx == 0:
                set_cell_shading(cell, PALE_BLUE)
                for run in p.runs:
                    set_run(run, size=font_size, color=NAVY, bold=True)
    if header:
        set_repeat_table_header(table.rows[0])
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    return table


def add_hyperlink(paragraph, text, url):
    rel = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    r_pr.extend([fonts, color, underline, size])
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def integration_doc():
    doc = Document()
    configure_page(doc)
    configure_styles(doc, "compact")
    nums = add_numbering(doc, "compact")
    set_header_footer(doc, "Пакет подключения", "T-Банк + СДЭК")
    add_title_block(
        doc,
        "Пакет подключения для заказчика",
        "Оплата и доставка Regola",
        "Что уже готово, какие данные нужны и как безопасно запустить боевой контур",
        ("Regola / regola.shop", "Готово к тестовым ключам", "14 августа 2026"),
    )
    add_callout(doc, "Ключевой вывод", "Код интеграции готов. Реальные платежи намеренно не включаются без договора, ключей терминала, налоговых настроек и приёмочного теста заказчика.")
    doc.add_paragraph("1. Что уже реализовано", style="Heading 1")
    add_list(doc, [
        "Сервер сам загружает товары и цены из базы; цена и итог из браузера игнорируются.",
        "T-Банк: Init, чек, Token, платёжная ссылка, success/fail возвраты и проверяемый webhook.",
        "СДЭК API v2: OAuth, город, тариф, ПВЗ, стоимость и создание отправления из админки.",
        "Админка показывает статус платежа, PaymentId, номер и UUID отправления СДЭК.",
        "Без ключей сайт остаётся рабочим: принимает заказ по счёту, а онлайн-оплата скрыта.",
    ], nums["bullet"], "compact")
    doc.add_paragraph("2. Граница ответственности", style="Heading 2")
    add_paragraph(doc, "СДЭК отвечает за доставку и, при отдельном условии договора, может принять наложенный платёж. Онлайн-эквайринг банковской картой выполняет T-Банк. Это два независимых договора и два набора ключей.")

    page_break(doc)
    doc.add_paragraph("T-Банк: что получить у клиента", style="Heading 1")
    add_callout(doc, "Передача секретов", "TerminalKey можно идентифицировать в переписке, но пароль терминала передаётся только защищённым каналом и хранится исключительно в .env на VPS.", "warning")
    add_table(doc, [
        ["Данные", "Что должен подтвердить клиент"],
        ["Договор", "Интернет-эквайринг активирован для ИП/ООО и домена regola.shop"],
        ["Доступ", "Тестовые и боевые TerminalKey + пароль терминала"],
        ["Налоги", "СНО; НДС товара; НДС доставки - письменно от бухгалтера"],
        ["Чеки", "Кто фискализирует: T-Банк, облачная или собственная касса"],
        ["Сценарий", "Одностадийная оплата; разрешённые способы на форме банка"],
        ["Возвраты", "Ответственный сотрудник и регламент полной/частичной отмены"],
    ], [2700, 6660])
    doc.add_paragraph("Переменные сервера", style="Heading 2")
    add_list(doc, [
        "TBANK_MODE=test",
        "TBANK_TERMINAL_KEY и TBANK_PASSWORD",
        "TBANK_TAXATION",
        "TBANK_ITEM_TAX и TBANK_DELIVERY_TAX",
        "PUBLIC_BASE_URL=https://regola.shop",
    ], nums["bullet"], "compact")
    add_paragraph(doc, "Адрес уведомлений: https://regola.shop/api/payments/tbank/notification", bold_lead="Адрес уведомлений:")

    page_break(doc)
    doc.add_paragraph("СДЭК: что получить у клиента", style="Heading 1")
    add_table(doc, [
        ["Данные", "Что должен подтвердить клиент"],
        ["Договор", "Договор СДЭК для ИП/ООО и доступ к API v2"],
        ["Наложенный платёж", "Подключён отдельно; известны комиссия, лимиты и срок перечисления"],
        ["Ключи", "Тестовые и боевые Client ID + Client secret"],
        ["Отправитель", "Код города, shipment point, название, телефон"],
        ["Тарифы", "Разрешённые коды тарифов до ПВЗ и курьером"],
        ["Упаковка", "Вес в граммах и длина/ширина/высота в сантиметрах"],
    ], [2700, 6660])
    doc.add_paragraph("Переменные сервера", style="Heading 2")
    add_list(doc, [
        "CDEK_MODE=test; CDEK_CLIENT_ID; CDEK_CLIENT_SECRET",
        "CDEK_FROM_CITY_CODE; CDEK_TARIFF_PVZ; CDEK_TARIFF_COURIER",
        "CDEK_PACKAGE_WEIGHT_G; CDEK_PACKAGE_LENGTH_CM; CDEK_PACKAGE_WIDTH_CM; CDEK_PACKAGE_HEIGHT_CM",
        "CDEK_SHIPMENT_POINT; CDEK_SENDER_NAME; CDEK_SENDER_PHONE",
    ], nums["bullet"], "compact")
    add_callout(doc, "Важно", "Вес и габариты сейчас задаются как параметры одной единицы. Если упаковка сильно отличается по товарам, следующим этапом нужно хранить логистические параметры в каждой карточке.")

    page_break(doc)
    doc.add_paragraph("Приёмка и безопасный запуск", style="Heading 1")
    add_list(doc, [
        "Внести тестовые ключи на VPS и оставить оба режима test.",
        "Провести успешный платёж, отказ, отмену, повтор webhook, неверную сумму и возврат.",
        "Проверить чек: позиции, доставка, СНО, НДС, email/телефон покупателя.",
        "Проверить ПВЗ, курьера, предоплату, наложенный платёж и ошибочный город.",
        "Сверить платежи, чеки и отправления в кабинетах T-Банка и СДЭК.",
        "Получить письменное принятие заказчика, заменить ключи на боевые и включить production.",
        "Сделать один минимальный реальный заказ и только после сверки открыть оплату всем клиентам.",
    ], nums["number"], "compact")
    doc.add_paragraph("Тестовая матрица", style="Heading 2")
    add_table(doc, [
        ["Контур", "Проверка", "Ожидаемый результат"],
        ["T-Банк", "Успех / отказ / отмена", "Статус заказа соответствует уведомлению"],
        ["T-Банк", "Подмена Amount/PaymentId", "Webhook отклонён, статус не меняется"],
        ["СДЭК", "ПВЗ / курьер", "Цена, срок и код тарифа сохранены"],
        ["СДЭК", "Предоплата / наложенный", "Payment в накладной соответствует сценарию"],
        ["Админка", "Создание отправления дважды", "Повтор не создаёт дубликат"],
    ], [1800, 3000, 4560], font_size=9)

    page_break(doc)
    doc.add_paragraph("Форма передачи данных", style="Heading 1")
    add_paragraph(doc, "Заполняет заказчик или его бухгалтер/логист. Не вставлять секреты в этот документ - передать их отдельно.")
    add_table(doc, [
        ["Поле", "Ответ клиента"],
        ["Юридическое лицо / ИП, ИНН, ОГРН", ""],
        ["Система налогообложения", ""],
        ["НДС товара / доставки", ""],
        ["Касса и фискализация", ""],
        ["Кто оплачивает доставку", ""],
        ["Разрешён ли наложенный платёж", ""],
        ["Город и точка отгрузки", ""],
        ["Тарифы ПВЗ / курьер", ""],
        ["Вес и габариты упаковки", ""],
        ["Ответственный за приёмку", ""],
    ], [3300, 6060])
    doc.add_paragraph("Официальные источники", style="Heading 2")
    sources = [
        ("T-Банк: OpenAPI и адреса сред", "https://developer.tbank.ru/eacq/intro/developer/openapi"),
        ("T-Банк: формирование Token", "https://developer.tbank.ru/eacq/intro/developer/token"),
        ("T-Банк: метод Init", "https://developer.tbank.ru/eacq/api/init"),
        ("T-Банк: уведомления", "https://developer.tbank.ru/eacq/intro/developer/notification"),
        ("СДЭК: документация API", "https://api-docs.cdek.ru/"),
    ]
    for label, url in sources:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_hyperlink(p, label, url)

    configure_page(doc)
    path = OUT / "Regola_подключение_TBank_CDEK.docx"
    doc.save(path)
    return path


def design_doc():
    doc = Document()
    configure_page(doc)
    configure_styles(doc, "narrative")
    nums = add_numbering(doc, "narrative")
    set_header_footer(doc, "Дизайн-концепция", "Regola 2.0")
    add_editorial_cover(
        doc,
        "Концепция и техническое задание",
        "REGOLA / QUIET PRECISION",
        "Новый цифровой образ бренда дверной фурнитуры",
        "Версия для согласования | 14 августа 2026",
    )

    page_break(doc)
    doc.add_paragraph("1. Задача редизайна", style="Heading 1")
    add_paragraph(doc, "Сайт должен перестать выглядеть как витрина маркетплейса и стать самостоятельным брендовым инструментом: показывать качество металла, объяснять выбор, помогать дизайнеру сравнить отделки и уверенно доводить покупателя до заказа.")
    add_callout(doc, "Предлагаемое направление", "Quiet Precision / Тихая точность: архитектурная сетка, крупная предметная съёмка, спокойный тёмно-синий фон, тёплый металл как акцент, минимум декоративного шума.")
    doc.add_paragraph("Цели", style="Heading 2")
    add_list(doc, [
        "Повысить доверие к Regola как к бренду, а не только продавцу на маркетплейсах.",
        "Сделать каталог первым и самым сильным сценарием входа.",
        "Упростить подбор по типу, цвету, форме, механизму и стилю интерьера.",
        "Дать безошибочное оформление заказа на мобильном устройстве.",
        "Сохранить текущую сине-песочную гамму, но сделать её глубже и премиальнее.",
    ], nums["bullet"], "narrative")
    doc.add_paragraph("Аудитории", style="Heading 2")
    add_table(doc, [
        ["Сегмент", "Главный вопрос", "Что должен дать интерфейс"],
        ["Владелец жилья", "Подойдёт ли к интерьеру и двери?", "Фото, отделки, размеры, понятный заказ"],
        ["Дизайнер / архитектор", "Можно ли уверенно заложить в проект?", "Фильтры, спецификация, материалы, контакты"],
        ["Комплектатор / подрядчик", "Есть ли стабильность и понятная поставка?", "Артикулы, сроки, оптовый контакт, документы"],
    ], [2000, 3200, 4160], font_size=9)

    page_break(doc)
    doc.add_paragraph("2. Референсы: принципы, не копирование", style="Heading 1")
    add_paragraph(doc, "Референсы задают уровень арт-дирекции и структуру опыта. Их композиции, тексты, фотографии и фирменные элементы не копируются.")
    add_table(doc, [
        ["Бренд", "Что изучаем", "Как применяем в Regola"],
        ["FORMANI", "Коллекции как система, авторство, спокойный ритм", "Разделять продукт, коллекцию и историю; больше воздуха"],
        ["Olivari", "Каталог, типология, дизайнеры, техническая ясность", "Фильтры, связанная коллекция, характеристики"],
        ["Buster + Punch", "Тактильные макро, отделки, уверенный контраст", "Крупный металл, честная фактура, акцент на покрытии"],
    ], [1800, 3500, 4060], font_size=9)
    doc.add_paragraph("Ссылки", style="Heading 2")
    for label, url in [
        ("FORMANI - официальный сайт", "https://formani.com/"),
        ("Olivari - официальный каталог ручек", "https://www.olivari.it/en/door-handles/"),
        ("Buster + Punch - door hardware", "https://busterandpunch.com/product-category/door-hardware/"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_hyperlink(p, label, url)
    add_callout(doc, "Антиреференс", "Шаблонный интернет-магазин с мелким логотипом, перегруженным меню, одинаковыми белыми карточками и набором случайных иконок. Новый дизайн должен быть характерным даже без логотипа.", "warning")

    page_break(doc)
    doc.add_paragraph("3. Визуальная система", style="Heading 1")
    doc.add_paragraph("Цвет", style="Heading 2")
    add_table(doc, [
        ["Токен", "HEX", "Назначение"],
        ["Regola Navy", "#0E2238", "Навигация, крупные фоны, premium-контраст"],
        ["Deep Ink", "#071522", "Hover, футер, акцентные экраны"],
        ["Warm Ivory", "#F4F0E8", "Основной светлый фон"],
        ["Sand", "#D8B68A", "Мягкие разделители и вторичный акцент"],
        ["Bronze", "#A86F42", "CTA и микроакценты, дозированно"],
        ["Graphite", "#27313A", "Основной текст"],
    ], [2500, 1600, 5260])
    doc.add_paragraph("Типографика", style="Heading 2")
    add_list(doc, [
        "Заголовки: Manrope 600/700 - современно, геометрично, с хорошей кириллицей.",
        "Текст и интерфейс: Inter 400/500/600 - нейтрально и очень читаемо.",
        "Шкала desktop: 64/72, 44/52, 30/38, 20/28, 16/26, 14/20.",
        "На mobile крупный заголовок уменьшается до 38/44; основной текст остаётся не меньше 16 px.",
    ], nums["bullet"], "narrative")
    doc.add_paragraph("Сетка и ритм", style="Heading 2")
    add_paragraph(doc, "Desktop: контейнер 1280 px, 12 колонок, поля 32-64 px. Tablet: 8 колонок. Mobile: 4 колонки, поля 16 px. Базовый шаг 8 px; секции 96-144 px на desktop и 64-88 px на mobile.")

    page_break(doc)
    doc.add_paragraph("4. Архитектура и навигация", style="Heading 1")
    add_table(doc, [
        ["Уровень", "Разделы"],
        ["Основная навигация", "Каталог / О бренде / Как заказать / Гарантии / Контакты"],
        ["Утилиты", "Поиск / Корзина / Выбор города или доставки"],
        ["Каталог", "Все ручки / На розетке / Кнобы / Комплекты / Новинки"],
        ["Фильтры", "Тип / форма / цвет / покрытие / механизм / цена"],
        ["Служебные", "Оплата и доставка / возврат / реквизиты / 404"],
    ], [2300, 7060])
    doc.add_paragraph("Главная: порядок блоков", style="Heading 2")
    add_list(doc, [
        "Первый экран: сильная предметная фотография + короткое обещание + CTA Каталог.",
        "Каталог сразу под первым экраном: поиск, фильтры и 6-8 ключевых моделей.",
        "Выбор по отделке: визуальные swatches, крупные макро деталей.",
        "Почему Regola: технология, проверка качества, материалы - без общих рекламных фраз.",
        "Как заказать: 3 ясных шага, T-Банк и СДЭК как сервисные преимущества.",
        "Производство и контроль качества; затем контакты и полноценный футер.",
    ], nums["number"], "narrative")
    add_callout(doc, "Навигация", "На desktop меню находится в одной ясной строке под/рядом с увеличенным логотипом. На mobile - компактная шапка, поиск и полноэкранное меню; нижней закреплённой контактной плашки нет.")

    page_break(doc)
    doc.add_paragraph("5. Ключевые страницы", style="Heading 1")
    doc.add_paragraph("Каталог", style="Heading 2")
    add_list(doc, [
        "Заголовок и краткий результат: количество моделей, активные фильтры, сортировка.",
        "Карточка: фото 4:5 без обрезки изделия, название, цена от, отделка, быстрый просмотр.",
        "Фильтры на desktop слева, на mobile в bottom sheet; активные фильтры видны chips.",
        "Поиск с подсказками по названию, артикулу, цвету и типу.",
        "Скелетоны загрузки, пустая выдача с понятным сбросом, стабильная сетка без скачков.",
    ], nums["bullet"], "narrative")
    doc.add_paragraph("Карточка товара", style="Heading 2")
    add_list(doc, [
        "Галерея 1-6 фото, превью, zoom, swipe и видео при наличии.",
        "Название, цена от, артикул, краткая ценность, отделки и количество для заказа.",
        "CTA Добавить в корзину; ссылки на Wildberries, Ozon и Яндекс Маркет как альтернативные каналы.",
        "Характеристики, комплектация, совместимость, гарантия, доставка и FAQ.",
        "Связанные модели/отделки и просмотренные товары.",
    ], nums["bullet"], "narrative")
    doc.add_paragraph("Корзина и checkout", style="Heading 2")
    add_paragraph(doc, "Один линейный сценарий: контакты -> город -> тариф/ПВЗ -> способ оплаты -> проверка состава -> переход в T-Банк. Итог и доставка всегда пересчитываются сервером. Ошибка должна объяснять, что исправить, не очищая введённые данные.")

    page_break(doc)
    doc.add_paragraph("6. Компоненты и состояния", style="Heading 1")
    add_table(doc, [
        ["Компонент", "Обязательные состояния"],
        ["Кнопка", "default / hover / pressed / focus / disabled / loading"],
        ["Поле", "empty / filled / focus / valid / error / disabled"],
        ["Карточка", "default / hover / quick view / unavailable / loading"],
        ["Фильтр", "default / selected / count / reset / no results"],
        ["Платёж", "creating / redirect / pending / paid / failed / refunded"],
        ["Доставка", "calculating / PVZ list / selected / API error / manual fallback"],
    ], [2500, 6860])
    doc.add_paragraph("Motion", style="Heading 2")
    add_list(doc, [
        "120-180 ms для hover/focus; 220-320 ms для раскрытия и перехода между состояниями.",
        "Плавный скролл к якорям только при явном действии пользователя.",
        "Галерея не двигает разметку; все медиа резервируют размер заранее.",
        "Уважать prefers-reduced-motion: отключать параллакс и сложные появления.",
        "Никаких бесконечных декоративных анимаций и тяжёлого фонового видео на mobile.",
    ], nums["bullet"], "narrative")
    doc.add_paragraph("Фотостиль", style="Heading 2")
    add_paragraph(doc, "70% - изделие на нейтральном тёплом фоне; 20% - макро покрытия и механизма; 10% - интерьерный контекст. Один свет, одна температура, одинаковый масштаб. Фото не должны содержать текст маркетплейсов, плашки и чужие водяные знаки.")

    page_break(doc)
    doc.add_paragraph("7. Mobile-first требования", style="Heading 1")
    add_list(doc, [
        "Контрольные ширины: 360, 390, 430, 768, 1024, 1280 и 1440 px.",
        "Интерактивная зона не меньше 44 x 44 px; между соседними целями достаточно воздуха.",
        "Никакого горизонтального скролла, обрезанных заголовков и таблиц шире экрана.",
        "Фильтры открываются снизу, сохраняют выбор и показывают количество результатов.",
        "Checkout использует нативные типы полей, автозаполнение, цифровую клавиатуру телефона.",
        "Выбор ПВЗ читаем без карты: адрес, режим работы, код; карта может быть вторым этапом.",
        "Корзина и CTA доступны большим пальцем, но не перекрывают контент и футер.",
    ], nums["bullet"], "narrative")
    doc.add_paragraph("Доступность", style="Heading 2")
    add_list(doc, [
        "WCAG 2.2 AA: контраст текста, видимый focus, управление клавиатурой.",
        "Корректная семантика headings/landmarks, alt для предметных фото, labels для форм.",
        "Ошибки формы связаны с полем и озвучиваются assistive technology.",
        "Цвет не является единственным способом передать статус или выбор.",
    ], nums["bullet"], "narrative")

    page_break(doc)
    doc.add_paragraph("8. Производительность, SEO и аналитика", style="Heading 1")
    add_table(doc, [
        ["Область", "Критерий приёмки"],
        ["Core Web Vitals", "LCP <= 2.5 s, CLS <= 0.1, INP <= 200 ms на типовом mobile"],
        ["Изображения", "AVIF/WebP, srcset, lazy-load ниже fold, явные width/height"],
        ["JS/CSS", "Минимум зависимостей, code splitting, без блокирующих виджетов"],
        ["SEO", "Уникальные title/description, canonical, sitemap, robots, Product schema"],
        ["Аналитика", "Поиск, фильтр, просмотр товара, add-to-cart, checkout, payment result"],
        ["Ошибки", "Логи payment/CDEK без секретов и персональных данных"],
    ], [2500, 6860])
    add_callout(doc, "Бюджет качества", "Фото не должны разрушать скорость. Визуальная премиальность достигается качеством кадра, сеткой и типографикой, а не тяжёлыми эффектами.")
    doc.add_paragraph("Контент, который нужен от клиента", style="Heading 2")
    add_list(doc, [
        "Логотип в SVG, правила использования и подтверждённая цветовая гамма.",
        "Категории, артикула, размеры, материалы, покрытия, комплектация, гарантия.",
        "Фотографии без плашек маркетплейсов: минимум 2000 px по длинной стороне.",
        "Юридические реквизиты, условия оплаты/доставки/возврата и контакты.",
        "Ответы на вопросы о производстве, контроле качества и происхождении материалов.",
    ], nums["bullet"], "narrative")

    page_break(doc)
    doc.add_paragraph("9. Этапы и результат", style="Heading 1")
    add_table(doc, [
        ["Этап", "Результат", "Точка согласования"],
        ["1. Discovery", "Контент, аудит, аудитории, ограничения", "Направление и приоритеты"],
        ["2. UX", "Карта сайта, пользовательские потоки, wireframes", "Структура ключевых страниц"],
        ["3. Visual", "Moodboard, 2 направления, UI-kit", "Один утверждённый концепт"],
        ["4. Design", "Desktop + mobile макеты, прототип", "Полный дизайн"],
        ["5. Build", "Адаптивная реализация и миграция контента", "Staging"],
        ["6. QA", "Функционал, device, a11y, speed, SEO", "Разрешение на production"],
    ], [1500, 4260, 3600], font_size=9)
    doc.add_paragraph("Критерии готовности", style="Heading 2")
    add_list(doc, [
        "Все страницы и состояния утверждены для desktop и mobile.",
        "Нет горизонтального скролла, layout shift, тупиковых сценариев и непонятных ошибок.",
        "Каталог, поиск, карточка, корзина, T-Банк, СДЭК и админка проходят приёмочные тесты.",
        "Контент и юридические тексты предоставлены и утверждены клиентом.",
        "Метрики производительности и доступности достигнуты на staging.",
        "Есть резервная копия, план отката и список переменных production.",
    ], nums["bullet"], "narrative")

    page_break(doc)
    doc.add_paragraph("10. Что согласовать до начала дизайна", style="Heading 1")
    add_table(doc, [
        ["Решение", "Ответ клиента"],
        ["Подтверждаем направление Quiet Precision", "Да / Нет / Комментарий"],
        ["Какой сегмент главный: B2C, дизайнеры или комплектация", ""],
        ["Какие категории запускаем в первой версии", ""],
        ["Можно ли заказать напрямую или только через менеджера/маркетплейс", ""],
        ["Нужны ли оптовые цены и кабинет партнёра", ""],
        ["Есть ли качественная съёмка или нужна новая", ""],
        ["Кто утверждает тексты, фото, оплату и юридические условия", ""],
        ["Желаемый срок и бюджет редизайна", ""],
    ], [5800, 3560], font_size=9)
    add_callout(doc, "Рекомендация", "Сначала согласовать направление, контент и 2-3 ключевых экрана. Не переносить весь сайт в новый стиль до утверждения главной, каталога и карточки товара.")
    doc.add_paragraph("Итоговые артефакты следующего этапа", style="Heading 2")
    add_list(doc, [
        "Moodboard и 2 визуальные концепции главного экрана.",
        "UX-wireframes главной, каталога, карточки и checkout.",
        "UI-kit с токенами, компонентами и состояниями.",
        "Кликабельный mobile/desktop прототип.",
        "Передача разработке с спецификацией и контрольным дизайн-QA.",
    ], nums["bullet"], "narrative")

    configure_page(doc)
    path = OUT / "Regola_дизайн-концепция_и_ТЗ.docx"
    doc.save(path)
    return path


def audit_docx(path):
    doc = Document(path)
    for section in doc.sections:
        assert round(section.page_width.inches, 2) == 8.5
        assert round(section.page_height.inches, 2) == 11.0
        assert round(section.left_margin.inches, 2) == 1.0
        assert round(section.right_margin.inches, 2) == 1.0
    for table in doc.tables:
        grid = table._tbl.tblGrid
        widths = [int(col.get(qn("w:w"))) for col in grid]
        assert sum(widths) == TABLE_WIDTH_DXA, (path, widths)
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        assert tbl_w is not None and int(tbl_w.get(qn("w:w"))) == TABLE_WIDTH_DXA
        assert tbl_ind is not None and int(tbl_ind.get(qn("w:w"))) == TABLE_INDENT_DXA
        for row in table.rows:
            cell_widths = [int(cell._tc.tcPr.find(qn("w:tcW")).get(qn("w:w"))) for cell in row.cells]
            assert cell_widths == widths, (path, cell_widths, widths)
    assert not any(p.text.strip().startswith(("- ", "* ")) for p in doc.paragraphs)


if __name__ == "__main__":
    paths = [integration_doc(), design_doc()]
    for item in paths:
        audit_docx(item)
        print(item)
